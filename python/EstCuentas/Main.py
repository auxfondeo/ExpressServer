# -*- coding: utf-8 -*-

import win32com.client
import win32com.client as win32
import datetime
import re,os
import openpyxl
from openpyxl import Workbook
from datetime import datetime
import openpyxl
from openpyxl import Workbook
from datetime import datetime
from openpyxl.utils import get_column_letter
from openpyxl import load_workbook
import zipfile,sys

##Funcion que reemplaza el texto
def reemplazar_texto(texto_interno, original, reemplazo):
    return texto_interno.replace(original, reemplazo)

##Funcion que remplaza los espacios en exceso del texto
def corregir_texto(texto):
    texto_corregido = " ".join(texto.split())
    return texto_corregido

##Funcion que corrige el texto
def corregir_textos(texto):
    texto_corregido = texto.rstrip()
    return texto_corregido


##Funcion que elimina los acentos del texto
def remover_acentos(texto):
    texto_sin_acentos = texto
    reemplazos = [
        {"Á": "A"}, {"É": "E"}, {"Í": "I"}, {"Ó": "O"}, {"Ú": "U"},
        {"À": "A"}, {"È": "E"}, {"Ì": "I"}, {"Ò": "O"}, {"Ù": "U"},
        {"Ä": "A"}, {"Ë": "E"}, {"Ï": "I"}, {"Ö": "O"}, {"Ü": "U"},
        {"á": "a"}, {"é": "e"}, {"í": "i"}, {"ó": "o"}, {"ú": "u"},
        {"à": "a"}, {"è": "e"}, {"ì": "i"}, {"ò": "o"}, {"ù": "u"},
        {"ä": "a"}, {"ë": "e"}, {"ï": "i"}, {"ö": "o"}, {"ü": "u"}
    ]
    for replacement in reemplazos:
        original = list(replacement.keys())[0]
        reemplazo = list(replacement.values())[0]
        texto_sin_acentos = reemplazar_texto(texto_sin_acentos, original, reemplazo)
    texto_sin_acentos = corregir_texto(texto_sin_acentos)
    textos_corregidos = corregir_textos(texto_sin_acentos)

    return textos_corregidos

##Funcion que crea Estado de cuenta individual
def EstadoCuenta(numero_credito,nombre):
    import os
    #print("Segundo MARCADOR")
    
    # Este script automatiza el envío de estados de cuenta por correo electrónico.
    # Importar las librerías necesarias
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import pandas as pd
    import time
    from selenium.webdriver.common.by import By
    import os
    from datetime import datetime
    directorio_actual = os.path.dirname(os.path.abspath(__file__))
    print(directorio_actual)
    
    texto = nombre
    Cliente = remover_acentos(texto)
    print("Nombre",Cliente)
    
    buscador = numero_credito


# Eliminamos los archivos 'pagosCredito.xlsx', 'movimientosCredito.xlsx', 'formatoEstadoCuenta_filled.docx' combined_data.xlsx y  si existen
    file_path = os.path.join(directorio_actual, "pagosCredito.xlsx")
    #file_path = 'C:/Users/alber/Downloads/pagosCredito.xlsx'
    
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'pagosCredito.xlsx' eliminado.")
    else:
        print()
        #print("El archivo 'pagosCredito.xlsx' no existe.")
    file_path = os.path.join(directorio_actual, "movimientosCredito.xlsx")
    #file_path = 'C:/Users/alber/Downloads/movimientosCredito.xlsx'
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'movimientosCredito.xlsx' eliminado.")
    else:
        print()
        #print("El archivo 'movimientosCredito.xlsx' no existe.")

    file_path = os.path.join(directorio_actual, "formatoEstadoCuenta_filled.xlsx")
    #file_path = 'C:/Users/alber/Downloads/formatoEstadoCuenta_filled.docx'
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'formatoEstadoCuenta_filled.docx' eliminado.")
    else:
        print()
        #print("El archivo 'formatoEstadoCuenta_filled.docx' no existe.")

    file_path = os.path.join(directorio_actual, "combined_data.xlsx")
    #file_path = 'C:/Users/alber/Downloads/combined_data.xlsx'
    
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'combined_data.xlsx' eliminado.")
    else:
        print()
        #print("El archivo 'combined_data.xlsx' no existe.")

    file_path = os.path.join(directorio_actual, "EstadoCuenta.pdf")
    #file_path = 'C:/Users/alber/Downloads/EstadoCuenta.pdf'
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'formatoEstadoCuenta.pdf' eliminado.")
    else:
        print()
        #print("El archivo 'formatoEstadoCuenta.pdf' no existe.")

# Configurar el servicio del controlador de Chrome
    s = Service()  # O "C:/Drivers/chromedriver.exe"

# Inicializar el controlador de Chrome
    driver = webdriver.Chrome(service=s)

# Abrir la página de inicio de sesión
    driver.get("https://valora.credisoft4.com/Acceso/Login")
    driver.maximize_window()

    time.sleep(1)  # Esperar 1 segundo

# Rellenar el usuario y contraseña, luego presionar ENTER
    usuario = driver.find_element(By.XPATH, "//input[contains(@id,'Usuario')]")
    usuario.send_keys("MARCO.OVIEDO" + Keys.TAB+ "Valora24" + Keys.ENTER)

# Esperar a que la página cargue y el enlace sea clickeable
    WebDriverWait(driver, 10).until(
    EC.element_to_be_clickable((By.XPATH, "//a[@href='/Reporte']"))
)

# Clickear en el enlace de reportes
    reportes = driver.find_element(By.XPATH, "//span[@class='menu-text'][contains(.,'Créditos')]")
    reportes.click()
    time.sleep(1)  # Esperar 1 segundo

# Clickear en el botón de listado
    listado = driver.find_element(By.XPATH, "//a[@href='/Credito']")
    listado.click()
    time.sleep(1)  # Esperar 1 segundo

# Rellenar la búsqueda avanzada
    busqueda = driver.find_element(By.XPATH, "//input[@type='search']")
    busqueda.send_keys(buscador + Keys.ENTER)

# Esperar a que desaparezca el elemento "Procesando..."
    WebDriverWait(driver, 100).until(
    EC.invisibility_of_element_located((By.ID, "TablaCreditos_processing"))
)
    
    time.sleep(3)

# Después de rellenar la búsqueda avanzada y esperar a que la página se cargue
    html = driver.page_source
    soup = BeautifulSoup(html, 'html.parser')

# Buscar la tabla por su id
    tabla = soup.find('table', id='TablaCreditos')

# Extraer los encabezados de la tabla
    encabezados = [th.get_text().strip() for th in tabla.find_all('th')]

# Inicializar una lista para almacenar los datos de la tabla
    datos_tabla = []

# Iterar sobre cada fila de la tabla
    for fila in tabla.find('tbody').find_all('tr'):
        datos_fila = [td.get_text().strip() for td in fila.find_all('td')]
        datos_tabla.append(datos_fila)

# Convertir la lista de datos en un DataFrame
    datos_cliente = pd.DataFrame(datos_tabla, columns=encabezados)
    #print(datos_cliente)

    valor_deseado = Cliente
    #print("valor deseado:"+valor_deseado)

    # Comparar la columna 'Cliente' con el valor deseado
    resultado_comparacion = datos_cliente['Cliente'] == valor_deseado
    if resultado_comparacion.all() == False:
        #print("Nombre no coincide")
        sys.exit()
# Clickear en el botón para abrir el número de crédito
    Detalles = driver.find_element(By.XPATH, "//i[contains(@class,'fa fa-bars fa-fw')]") 
    Detalles.click()
    time.sleep(1)  # Esperar 1 segundo

# Clickear en la pestaña de pagos
    Pagos = driver.find_element(By.XPATH, "//a[@data-toggle='tab'][contains(.,'Pagos')]")
    Pagos.click()
    time.sleep(1)  # Esperar 1 segundo

# Esperar a que la tabla se cargue completamente
    WebDriverWait(driver, 10).until(
    EC.presence_of_element_located((By.XPATH, "//table"))
)

# Extraer el código HTML de la página
    html = driver.page_source
    soup = BeautifulSoup(html, 'html.parser')

# Buscar los encabezados de la tabla
    table_headers = soup.find_all('th', class_='sorting_disabled')

# Extraer el texto de cada encabezado
    headers = [th.get_text().strip() for th in table_headers]

# Encontrar la tabla por su identificador único
    tabla = soup.find('table', id='example')

# Inicializar una lista para almacenar los datos de la tabla
    datos_tabla = []

# Iterar sobre cada fila de la tabla, excluyendo la primera fila que contiene los encabezados
    for fila in tabla.find_all('tr')[1:]:
        datos_fila = []  # Inicializar una lista para los datos de la fila actual
    # Encontrar todas las celdas en la fila
        celdas = fila.find_all('td')
    # Extraer el texto de cada celda y agregarlo a la lista de datos de la fila
        datos_fila.extend(celda.get_text().strip() for celda in celdas)
    # Agregar la fila de datos a la lista principal
        datos_tabla.append(datos_fila)

# Convertir la lista de listas en un DataFrame y asignar los encabezados apropiados
    df = pd.DataFrame(datos_tabla, columns=headers)

# Exportar el DataFrame a un archivo Excel en la carpeta de descargas
    ruta_destino_excel = os.path.join(directorio_actual, "pagosCredito.xlsx")
    # Guardar el DataFrame en el archivo de Excel
    df.to_excel(ruta_destino_excel, index=False)
    #df.to_excel(r"C:\Users\alber\Downloads\pagosCredito.xlsx", index=False)

# Clickear en la pestaña de movimientos
    Movimientos = driver.find_element(By.XPATH, "//a[@data-toggle='tab'][contains(text(),'Movimientos')]")
    Movimientos.click()
    time.sleep(1)  # Esperar 1 segundo

# Esperar a que la tabla se cargue completamente
    WebDriverWait(driver, 10).until(
    EC.presence_of_element_located((By.XPATH, "//table"))
)

# Extraer el código HTML de la página
    html = driver.page_source
    soup = BeautifulSoup(html, 'html.parser')

# Encontrar la tabla por su clase
    tabla = soup.find('table', class_='table table-striped table-condensed')

# Extraer los encabezados de la tabla
    headers = [th.get_text().strip() for th in tabla.find('thead').find_all('th')]
    headers = headers[:-1]  # Excluir el último encabezado que contiene el botón de eliminar

# Inicializar una lista para almacenar los datos de la tabla
    datos_tabla = []

# Iterar sobre cada fila de la tabla en el cuerpo de la tabla
    for fila in tabla.find('tbody').find_all('tr'):
        datos_fila = [td.get_text().strip() for td in fila.find_all('td')]
    # El último elemento 'td' contiene un enlace para eliminar, puedes decidir si incluirlo o no
        datos_tabla.append(datos_fila[:-1])  # Excluir el último 'td' que contiene el botón de eliminar

# Convertir la lista de listas en un DataFrame y asignar los encabezados apropiados
    df = pd.DataFrame(datos_tabla, columns=headers)

# Exportar el DataFrame a un archivo Excel en la carpeta de descargas
    # Ruta del archivo de destino (relativa al directorio actual)
    ruta_destino_excel = os.path.join(directorio_actual, "movimientosCredito.xlsx")

    # Guardar el DataFrame en el archivo de Excel
    df.to_excel(ruta_destino_excel, index=False)
    #df.to_excel(r"C:\Users\alber\Downloads\movimientosCredito.xlsx", index=False)

# Leer el archivo de pagosCredito.xlsx
    # Ruta del archivo de origen (relativa al directorio actual)
    ruta_archivo_excel = os.path.join(directorio_actual, "pagosCredito.xlsx")

    # Leer el archivo Excel en un DataFrame
    df_pagos = pd.read_excel(ruta_archivo_excel)
    #df_pagos = pd.read_excel(r"C:\Users\alber\Downloads\pagosCredito.xlsx")

# Extraer las columnas [1,2,3,4,5,19] y excluir la primera fila
    db_estadoCuenta = df_pagos.iloc[1:, [1, 2, 3, 4, 5, 19]].copy()

# Corregir los encabezados
    db_estadoCuenta.columns = ['NumeroCredito', 'FechaPago', 'Capital', 'Interes', 'IVA', 'Total']

# Eliminar las comas (",") de los números
    db_estadoCuenta['Capital'] = db_estadoCuenta['Capital'].str.replace(',', '')
    db_estadoCuenta['Interes'] = db_estadoCuenta['Interes'].str.replace(',', '')
    db_estadoCuenta['IVA'] = db_estadoCuenta['IVA'].str.replace(',', '')
    db_estadoCuenta['Total'] = db_estadoCuenta['Total'].str.replace(',', '')

# Eliminar los símbolos "$"
    db_estadoCuenta['Capital'] = db_estadoCuenta['Capital'].str.replace('$', '').astype(float)
    db_estadoCuenta['Interes'] = db_estadoCuenta['Interes'].str.replace('$', '').astype(float)
    db_estadoCuenta['IVA'] = db_estadoCuenta['IVA'].str.replace('$', '').astype(float)
    db_estadoCuenta['Total'] = db_estadoCuenta['Total'].str.replace('$', '').astype(float)

# Cambiar el tipo de dato
    db_estadoCuenta['NumeroCredito'] = db_estadoCuenta['NumeroCredito'].astype(int)
    db_estadoCuenta['FechaPago'] = pd.to_datetime(db_estadoCuenta['FechaPago'], dayfirst=True)

# Ordenar la tabla por 'FechaPago' y resetear el índice
    db_estadoCuenta = db_estadoCuenta.sort_values(by='FechaPago').reset_index(drop=False)
    db_estadoCuenta.rename(columns={'index': 'Indice'}, inplace=True)

# Leer el archivo de movimientosCredito.xlsx
# Ruta del archivo de origen (relativa al directorio actual)
    ruta_archivo_excel = os.path.join(directorio_actual, "movimientosCredito.xlsx")

    # Leer el archivo Excel en un DataFrame
    df_movimientos = pd.read_excel(ruta_archivo_excel)
    #df_movimientos = pd.read_excel(r"C:\Users\alber\Downloads\movimientosCredito.xlsx")

# Tomar en cuenta solo las columnas [1,4]
    db_movimientos = df_movimientos.iloc[:, [1, 4]].copy()

# Eliminar el prefijo "$"
    db_movimientos['Importe'] = db_movimientos['Importe'].str.replace(',', '')
    db_movimientos['Importe'] = db_movimientos['Importe'].str.replace('$', '').astype(float)

# Cambiar el tipo de dato de 'Fecha Aplicación' a fecha
    db_movimientos['Fecha Aplicación'] = pd.to_datetime(db_movimientos['Fecha Aplicación'], dayfirst=True)

# Ordenar la tabla por 'Fecha Aplicación' y resetear el índice
    db_movimientos = db_movimientos.sort_values(by='Fecha Aplicación').reset_index(drop=False)
    db_movimientos.rename(columns={'index': 'Indice'}, inplace=True)
    db_movimientos['Indice'] += 1  # Hacer que el índice comience en 1

# Combinar las tablas utilizando la nueva columna 'Indice'
    db_combinado = pd.merge(db_estadoCuenta, db_movimientos, on='Indice', how='left')

# Cerrar el navegador
    driver.quit()
    #print("Tercer marcador")
    from docx import Document
    import datetime
    import re

# Abrir el documento ("C:\Users\alber\Downloads\Valora Credit\Automatizaciones\Estados de Cuenta\formatoEstadoCuenta.docx")
    # Ruta del archivo de origen (relativa al directorio actual)
    ruta_archivo_docx = os.path.join(directorio_actual, "formatoEstadoCuenta.docx")
    #print("Tercer marcador 1/2")   
    # Abrir el documento en formato .docx
    doc = Document(ruta_archivo_docx)
    #doc = Document(r"C:\Users\alber\Downloads\formatoEstadoCuenta.docx")

# Obtener la fecha actual
    fecha_actual = datetime.date.today().strftime("%Y-%m-%d")
    
# Obtener el ID_PRESTAMO y NOMBRE del DataFrame datos_cliente
    ID_PRESTAMO = datos_cliente['No. Crédito'].values[0]
    NOMBRE = datos_cliente['Cliente'].values[0]
    plazo_texto = datos_cliente['Plazo'].values[0]
    plazo_numero = int(re.findall(r'\d+', plazo_texto)[0])
    pago = float(datos_cliente['Pago'].values[0].replace('$', '').replace(',', ''))
    SALDO_TOTAL = plazo_numero * pago
    db_combinado['Importe'] = db_combinado['Importe'].replace(r'[\$,]', '', regex=True).astype(float)
    Saldo_pendiente = SALDO_TOTAL - db_combinado['Importe'].fillna(0).sum()
    Fecha_Inicio = datos_cliente['Fecha Primer Pago'].values[0]

    from datetime import datetime

# Formatear fecha_actual y Fecha_Inicio a dd/mm/aaaa
    fecha_actual = datetime.now().strftime("%d/%m/%Y")

    try:
    # Intenta convertir si está en el formato 'yyyy-mm-dd'
        Fecha_Inicio = datetime.strptime(Fecha_Inicio, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
    # Si ya está en el formato correcto (o en otro formato), puedes dejarlo tal como está o manejarlo adecuadamente
        pass

# Formatear pago, SALDO_TOTAL y Saldo_pendiente a $#,##0.00;-$#,##0.00
    def format_currency(value):
        """
    Formats a numeric value as a currency string.
    Args:
        value (float): The numeric value to be formatted.
    Returns:
        str: The formatted currency string.
    Example:
        >>> format_currency(1000)
        '$1,000.00'
        >>> format_currency(-500)
        '-$500.00'
    """
        return "${:,.2f}".format(value) if value >= 0 else "-${:,.2f}".format(abs(value))

    pago = format_currency(pago)
    SALDO_TOTAL = format_currency(SALDO_TOTAL)
    Saldo_pendiente = format_currency(Saldo_pendiente)

# Definir los datos a reemplazar
    data = {
    "FECHA_ACTUAL": fecha_actual,
    "ID_PRESTAMO": ID_PRESTAMO,
    "NOMBRE": NOMBRE,
    "Fecha_Inicio": Fecha_Inicio,
    "DESCUENTO": pago,
    "PLAZO": plazo_texto,
    "SALDO_TOTAL": SALDO_TOTAL,
    "SALDO_PENDIENTE": Saldo_pendiente
}

# Reemplazar los placeholders en el documento
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
            # Convertir el valor a cadena antes de reemplazar
                paragraph.text = paragraph.text.replace("{{" + key + "}}", str(value))

# Remplazar los placeholders en las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                        # Convertir el valor a cadena antes de reemplazar
                            paragraph.text = paragraph.text.replace("{{" + key + "}}", str(value))

    from docx import Document
    import pandas as pd
    #print("cuarto marcador")
# Guardamos el datafram combinado en una nueva variable pero no tomamos en cuenta la columna 'Indice'
    db = db_combinado.iloc[:, 1:].copy()
    db = db.fillna("-")
    # Ruta del archivo de destino (relativa al directorio actual)
    ruta_archivo_destino = os.path.join(directorio_actual, "combined_data.xlsx")

    # Guardar el DataFrame en el archivo de Excel
    db.to_excel(ruta_archivo_destino, index=False)

    # Leer el archivo Excel en un DataFrame
    db = pd.read_excel(ruta_archivo_destino)

    #db.to_excel(r"C:\Users\alber\Downloads\combined_data.xlsx", index=False)
    #db = pd.read_excel(r"C:\Users\alber\Downloads\combined_data.xlsx") 

# A 'db' le cambiamos el formato de las fechas para que solo muestre en formato dd/mm/aaaa (esto porque el formato em lo da asi: 2022-01-01 00:00:00)

    db['FechaPago'] = db['FechaPago'].dt.strftime('%d/%m/%Y')
# Asegurarse de que 'Fecha Aplicación' es de tipo datetime
    db['Fecha Aplicación'] = pd.to_datetime(db['Fecha Aplicación'], dayfirst=True, errors='coerce')

# Ahora, aplicar el formateo
    db['Fecha Aplicación'] = db['Fecha Aplicación'].dt.strftime('%d/%m/%Y')

# Si la columna 'Fecha Aplicación' tiene valores nulos, rellenarlos con un guion
    db['Fecha Aplicación'] = db['Fecha Aplicación'].fillna("-")

# Le damos formato de pesos a las columnas 'Capital', 'Interes', 'IVA' y 'Total'
    db['Capital'] = db['Capital'].apply(lambda x: f"${float(x):,.2f}")
    db['Interes'] = db['Interes'].apply(lambda x: f"${float(x):,.2f}")
    db['IVA'] = db['IVA'].apply(lambda x: f"${float(x):,.2f}")
    db['Total'] = db['Total'].apply(lambda x: f"${float(x):,.2f}")
    db['Importe'] = db['Importe'].apply(lambda x: f"${float(x):,.2f}" if x != '-' else '-')

# Encontrar el marcador y añadir una tabla
    for paragraph in doc.paragraphs:
        if "{{INSERTAR_TABLA_OPERACIONES}}" in paragraph.text:
            # Crear una tabla justo después del marcador
            table = doc.add_table(rows=1, cols=len(db.columns))

            # Establecer un estilo de tabla (por ejemplo, 'Table Grid')
            table.style = 'Table Grid'

            # Agregar encabezados a la tabla
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(db.columns):
                hdr_cells[i].text = column_name

            # Añadir los datos del DataFrame a la tabla
            for _, row in db.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            # Eliminar el marcador de posición
            paragraph.text = ""

# Guardar el documento
    # Ruta del archivo de destino (relativa al directorio actual)
    ruta_archivo_destino = os.path.join(directorio_actual, "formatoEstadoCuenta_filled.docx")

    # Guardar el documento .docx
    doc.save(ruta_archivo_destino)
    #doc.save(r"C:\Users\alber\Downloads\formatoEstadoCuenta_filled.docx")

# Mandamos a imprimir el documento es decir lo convertimos en pdf
    from docx2pdf import convert
    nombre_archivo_destino_pdf = f"EstadoCuenta_{numero_credito}.pdf"
    ruta_archivo_destino_final = os.path.join(directorio_actual, nombre_archivo_destino_pdf)
    convert(ruta_archivo_destino, ruta_archivo_destino_final)
    
    ruta_del_archivo_generado = ruta_archivo_destino_final
    #ruta_del_archivo_generado = r"C:\Users\alber\Downloads\EstadoCuenta_"+ str(numero_credito)+".pdf"#,numero_credito
    #print(ruta_del_archivo_generado)
    file_path = 'C:/Users/alber/Downloads/pagosCredito.xlsx'
    #print(nombre_pdf)
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'pagosCredito.xlsx' eliminado.")
    else:
        print()
        #print("El archivo 'pagosCredito.xlsx' no existe.")

    file_path = 'C:/Users/alber/Downloads/movimientosCredito.xlsx'
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'movimientosCredito.xlsx' eliminado.")
    else:
        print()
        #print("El archivo 'movimientosCredito.xlsx' no existe.")

    file_path = 'C:/Users/alber/Downloads/formatoEstadoCuenta_filled.docx'
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'formatoEstadoCuenta_filled.docx' eliminado.")
    else:
        print()
        #print("El archivo 'formatoEstadoCuenta_filled.docx' no existe.")

    file_path = 'C:/Users/alber/Downloads/combined_data.xlsx'
    if os.path.exists(file_path):
        os.remove(file_path)
        #print("Archivo 'combined_data.xlsx' eliminado.")
    else:
        print()
        #print("El archivo 'combined_data.xlsx' no existe.")

    #print("Terminado...")
    #print("RUTA_ARCHIVO: " + ruta_del_archivo_generado)
    return ruta_del_archivo_generado

##Funcion que obtiene todos los estados de cuenta del cliente

def main(nombre):
    # Este script automatiza el envío de estados de cuenta por correo electrónico.
    # Importar las librerías necesarias
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from bs4 import BeautifulSoup
    import pandas as pd
    import time
    from selenium.webdriver.common.by import By
    import os
    from datetime import datetime
    
    texto = nombre
    print("Texto",texto)
    Cliente = remover_acentos(texto)
    print(Cliente)
    
    #buscador = nombre
    #buscador = "CORTEZ MARQUEZ CARLOS"
    numeros_creditos = []  # Lista para almacenar los números de crédito
    # Eliminamos los archivos 'pagosCredito.xlsx', 'movimientosCredito.xlsx', 'formatoEstadoCuenta_filled.docx' combined_data.xlsx y  si existen

    # Configurar el servicio del controlador de Chrome
    s = Service()  # O "C:/Drivers/chromedriver.exe"

    # Inicializar el controlador de Chrome
    driver = webdriver.Chrome(service=s)

    # Abrir la página de inicio de sesión
    driver.get("https://valora.credisoft4.com/Acceso/Login")
    driver.maximize_window()

    time.sleep(1)  # Esperar 1 segundo

    # Rellenar el usuario y contraseña, luego presionar ENTER
    usuario = driver.find_element(By.XPATH, "//input[contains(@id,'Usuario')]")
    usuario.send_keys("MARCO.OVIEDO" + Keys.TAB+ "Valora24" + Keys.ENTER)

    # Esperar a que la página cargue y el enlace sea clickeable
    WebDriverWait(driver, 10).until(
    EC.element_to_be_clickable((By.XPATH, "//a[@href='/Reporte']"))
    )

    # Clickear en el enlace de reportes
    reportes = driver.find_element(By.XPATH, "//span[@class='menu-text'][contains(.,'Créditos')]")
    reportes.click()
    time.sleep(1)  # Esperar 1 segundo

    # Clickear en el botón de listado
    listado = driver.find_element(By.XPATH, "//a[@href='/Credito']")
    listado.click()
    time.sleep(1)  # Esperar 1 segundo

    # Rellenar la búsqueda avanzada
    busqueda = driver.find_element(By.XPATH, "//input[@type='search']")
    busqueda.send_keys(Cliente + Keys.ENTER)

    # Esperar a que desaparezca el elemento "Procesando..."
    WebDriverWait(driver, 100).until(
    EC.invisibility_of_element_located((By.ID, "TablaCreditos_processing"))
    )
    
    time.sleep(5)

    # Obtener los elementos de la tabla
    tabla_creditos = driver.find_element(By.ID, "TablaCreditos")
    filas = tabla_creditos.find_elements(By.TAG_NAME, "tr")
    # Recorrer las filas e imprimir los datos

    for fila in filas[1:]:  # Empezamos desde la segunda fila para evitar la fila de encabezados
        datos = fila.find_elements(By.TAG_NAME, "td")
        if len(datos) < 2:
            req = False
            return req
        numero_credito = datos[0].text
        cliente = datos[1].text
        monto = datos[2].text
        plazo = datos[3].text
        pago = datos[4].text
        fecha_otorgamiento = datos[5].text
        fecha_primer_pago = datos[6].text
        fecha_ultimo_pago = datos[7].text
        fecha_ingreso = datos[8].text
        fecha_finiquito = datos[9].text
        estatus = datos[10].text

        #print("Número de Crédito:", numero_credito)

        #print("Cliente:", cliente)
        #print("Monto:", monto)
        #print("Plazo:", plazo)
        #print("Pago:", pago)
        #print("Fecha de Otorgamiento:", fecha_otorgamiento)
        #print("Fecha Primer Pago:", fecha_primer_pago)
        #print("Fecha Último Pago:", fecha_ultimo_pago)
        #print("Fecha Ingreso:", fecha_ingreso)
        #print("Fecha Finiquito:", fecha_finiquito) ⚠️⚠️⚠️⚠️⚠️
        #print("Estatus:", estatus)
        print()
        if estatus == 'Vigente':
            numeros_creditos.append(numero_credito)
    #return numeros_creditos,Cliente

    time.sleep(2)
    driver.quit()
    archivos_pdf = []
    print("Marcador")
    print(numeros_creditos)
    for credito in numeros_creditos:
        nombre_archivo_pdf = EstadoCuenta(credito,Cliente)
        #print(nombre_archivo_pdf)
        archivos_pdf.append(nombre_archivo_pdf)

    #print(archivos_pdf)
    #print(Cliente)

    nombre_zip = Cliente+".zip"
    directorio_actual = os.path.dirname(os.path.abspath(__file__))

    ruta_destino_zip = os.path.join(directorio_actual,nombre_zip)
    #print("Creando ZIP")
    
    # Crear un archivo ZIP en la ubicación especificada
    with zipfile.ZipFile(ruta_destino_zip, 'w') as zipf:
        # Agregar cada archivo PDF al archivo ZIP
        for archivo_pdf in archivos_pdf:
            # Extraer solo el nombre del archivo sin la ruta
            nombre_archivo = os.path.basename(archivo_pdf)
            # Agregar el archivo al ZIP
            zipf.write(archivo_pdf, arcname=nombre_archivo)

    # Imprimir la ruta del archivo ZIP generado
    print("RUTA_ARCHIVO: " + ruta_destino_zip)
    return ruta_destino_zip

#Creditos,Cliente = obtenerCuentas("ACUÑA AGUERO JORGE SAUL")
#print(Creditos)

#main("ACUÑA AGUERO JORGE SAUL")


if __name__ == "__main__":
     # Verifica si se proporciona el argumento del nombre del cliente
    if len(sys.argv) != 2:
        print("Uso: CartaReestructura.py <nombre_cliente>")
        sys.exit(1)

    # Obtiene el nombre del cliente del primer argumento
    nombre_cliente = sys.argv[1]
    #numero_credito = sys.argv[2]
    
    # Llama a la función principal con el nombre del cliente
    main(nombre_cliente)

 
